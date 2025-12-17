from fastapi import UploadFile, HTTPException
from typing import Optional
import mimetypes

class FileService:
    ALLOWED_EXTENSIONS = {
        'image': ['image/jpeg', 'image/png', 'image/gif', 'image/webp'],
        'pdf': ['application/pdf'],
        'word': [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @staticmethod
    async def validate_file(file: UploadFile) -> bool:
        """
        Validate uploaded file type and size.
        
        Args:
            file: Uploaded file
        
        Returns:
            True if valid
        
        Raises:
            HTTPException if invalid
        """
        # Check file size
        file_content = await file.read()
        await file.seek(0)  # Reset file pointer
        
        if len(file_content) > FileService.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File size exceeds maximum allowed size of {FileService.MAX_FILE_SIZE / (1024*1024)}MB"
            )
        
        # Check file type
        content_type = file.content_type
        all_allowed = (
            FileService.ALLOWED_EXTENSIONS['image'] +
            FileService.ALLOWED_EXTENSIONS['pdf'] +
            FileService.ALLOWED_EXTENSIONS['word']
        )
        
        if content_type not in all_allowed:
            raise HTTPException(
                status_code=400,
                detail=f"File type {content_type} not allowed. Allowed types: images, PDF, Word documents"
            )
        
        return True
    
    @staticmethod
    def get_file_category(content_type: str) -> str:
        """Determine file category from content type."""
        if content_type in FileService.ALLOWED_EXTENSIONS['image']:
            return 'image'
        elif content_type in FileService.ALLOWED_EXTENSIONS['pdf']:
            return 'pdf'
        elif content_type in FileService.ALLOWED_EXTENSIONS['word']:
            return 'word'
        return 'unknown'

    @staticmethod
    def generate_pdf(content: str, filename: str) -> str:
        """
        Generate a PDF file from text content.
        
        Args:
            content: The text content to write
            filename: The output filename
            
        Returns:
            Absolute path to the generated file
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import simpleSplit
            import os
            
            # Ensure output directory exists - save to temp dir or specific outputs dir
            output_dir = os.path.join(os.getcwd(), "outputs")
            os.makedirs(output_dir, exist_ok=True)
            
            filepath = os.path.join(output_dir, filename)
            
            c = canvas.Canvas(filepath, pagesize=letter)
            width, height = letter
            
            # Simple text wrapping
            text_object = c.beginText(40, height - 40)
            text_object.setFont("Helvetica", 12)
            
            lines = content.split('\n')
            for line in lines:
                # Wrap long lines
                wrapped_lines = simpleSplit(line, "Helvetica", 12, width - 80)
                for wrapped_line in wrapped_lines:
                    # Check if we need a new page
                    if text_object.getY() < 40:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(40, height - 40)
                        text_object.setFont("Helvetica", 12)
                    text_object.textLine(wrapped_line)
            
            c.drawText(text_object)
            c.save()
            
            return filepath
        except Exception as e:
            raise Exception(f"PDF generation error: {str(e)}")

    @staticmethod
    def generate_docx(content: str, filename: str, student_info: dict = None) -> str:
        """
        Generate a Word document from text content, parsing basic Markdown.
        Supports:
        - Student Header Block
        - Headers (#, ##, ###)
        - Tables (| col | col |)
        - Bold (**text**)
        - Lists (-, *)
        """
        try:
            import docx
            from docx.shared import Pt
            import os
            import re
            
            # Ensure output directory exists
            output_dir = os.path.join(os.getcwd(), "outputs")
            os.makedirs(output_dir, exist_ok=True)
            
            filepath = os.path.join(output_dir, filename)
            
            doc = docx.Document()
            
            # 1. Render Student Info Header if provided
            if student_info:
                # Add a few blank lines or just start
                # Format: Name: Value (Bold Label)
                
                header_data = [
                    ("Name", student_info.get("name")),
                    ("Matricule No", student_info.get("matricule")),
                    ("School", student_info.get("school")),
                    ("Department", student_info.get("department")),
                    ("Level", student_info.get("level"))
                ]
                
                for label, value in header_data:
                    if value:
                        p = doc.add_paragraph()
                        # Bold Label
                        run_label = p.add_run(f"{label}: ")
                        run_label.bold = True
                        # Normal Value
                        p.add_run(str(value))
                
                # Add separator line
                # doc.add_paragraph("_" * 40) # Removed per user request
                doc.add_paragraph("") # Spacing

            # Clean Content (Regex for LaTeX removal)
            # Remove $$ ... $$
            content = re.sub(r'\$\$(.*?)\$\$', r'\1', content, flags=re.DOTALL)
            # Remove \text{...}
            content = re.sub(r'\\text\{(.*?)\}', r'\1', content)
            # Replace \to with ->
            content = content.replace(r'\to', '->')
            # Remove \[ ... \]
            content = re.sub(r'\\\[(.*?)\\\]', r'\1', content, flags=re.DOTALL)
            
            lines = content.split('\n')
            
            # Helpers must be defined before use
            def process_bold(paragraph, text):
                """Helper to process **bold** text within a paragraph."""
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        clean_part = part[2:-2]
                        run = paragraph.add_run(clean_part)
                        run.bold = True
                    else:
                        paragraph.add_run(part)

            def render_table(doc_obj, rows):
                """Helper to render a collected markdown table."""
                if not rows: return
                header_row = [c.strip() for c in rows[0].split('|') if c.strip()]
                if not header_row: return

                table = doc_obj.add_table(rows=1, cols=len(header_row))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, text in enumerate(header_row):
                    hdr_cells[i].text = text
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
                
                for row_line in rows[1:]:
                    if '---' in row_line: continue
                    cols = [c.strip() for c in row_line.split('|') if c.strip()]
                    if len(cols) != len(header_row): continue
                    row_cells = table.add_row().cells
                    for i, text in enumerate(cols):
                        row_cells[i].text = text

            # Loop through lines
            in_code_block = False
            in_table = False
            table_lines = []

            for line in lines:
                # Code Block Detection (```)
                if '```' in line:
                    in_code_block = not in_code_block
                    continue 

                # Inside Code Block -> Render Monospace
                if in_code_block:
                    p = doc.add_paragraph()
                    run = p.add_run(line) 
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    continue

                stripped_line = line.strip()
                
                # Table Detection
                if stripped_line.startswith('|') and stripped_line.endswith('|'):
                    in_table = True
                    table_lines.append(stripped_line)
                    continue
                else:
                    if in_table:
                        # End of table detected
                        render_table(doc, table_lines)
                        in_table = False
                        table_lines = []
                
                # Skip empty lines if we just finished a table or generic spacing
                if not line.strip():
                    # Preserve some spacing but not excessive
                    # doc.add_paragraph("") 
                    continue

                # Header Detection
                if stripped_line.startswith('#'):
                    # Count #'s
                    level = len(stripped_line.split(' ')[0])
                    # Clean text
                    text = stripped_line.lstrip('#').strip()
                    if level <= 9:
                        heading = doc.add_heading(text, level=level)
                        # Force Black Color and Bold
                        from docx.shared import RGBColor
                        for run in heading.runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = True
                    else:
                        doc.add_paragraph(text)
                
                # List Detection
                elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
                    text = stripped_line[2:]
                    p = doc.add_paragraph(style='List Bullet')
                    process_bold(p, text)
                    
                # Normal Paragraph
                else:
                    p = doc.add_paragraph()
                    process_bold(p, stripped_line)
            
            # Flush any pending table at end of file
            if in_table:
                render_table(doc, table_lines)
            
            doc.save(filepath)
            
            return filepath
        except Exception as e:
            raise Exception(f"DOCX generation error: {str(e)}")

file_service = FileService()
